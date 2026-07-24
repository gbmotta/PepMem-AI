"""Exporta relatórios PepMem-AI em Markdown, DOCX e PDF.

Relatório simples, descritivo e organizado a partir da predição + narrativa.
Não altera números do modelo — só formata o que já foi calculado.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any


@dataclass
class ReportSection:
    title: str
    paragraphs: list[str] = field(default_factory=list)
    bullets: list[str] = field(default_factory=list)
    table_headers: list[str] | None = None
    table_rows: list[list[str]] | None = None


@dataclass
class PepMemReport:
    title: str
    subtitle: str
    generated_at: str
    sections: list[ReportSection] = field(default_factory=list)


def _band_label(prob: float) -> str:
    if prob >= 0.70:
        return "Alta confiança (≥ 70%) — priorizar ensaio"
    if prob >= 0.40:
        return "Intermediário (40–70%)"
    return "Baixa probabilidade (< 40%)"


def build_single_report(
    *,
    sequence: str,
    target_label: str,
    res: dict[str, Any],
    narrative: str | None = None,
    neighbors: list[dict[str, Any]] | None = None,
    shap_top: list[dict[str, Any]] | None = None,
    interval: str | None = None,
    in_project: bool = False,
) -> PepMemReport:
    """Monta relatório de predição única."""
    prob = float(res.get("pred_high_activity_prob") or 0)
    pmi = res.get("pmi")
    q = res.get("q_peptide")
    now = datetime.now().strftime("%d/%m/%Y %H:%M")

    report = PepMemReport(
        title="Relatório PepMem-AI",
        subtitle="Predição peptídeo × membrana · InovAI Lab / UFRN",
        generated_at=now,
    )

    report.sections.append(
        ReportSection(
            title="1. Identificação",
            bullets=[
                f"Sequência: {sequence}",
                f"Comprimento: {len(sequence)} aminoácidos",
                f"Membrana-alvo: {target_label}",
                f"No banco do projeto: {'sim' if in_project else 'não'}",
                f"Gerado em: {now}",
            ],
        )
    )

    metrics = [
        f"Probabilidade calibrada de alta atividade (MIC ≤ 3,4 µM): {prob:.1%}",
        f"Faixa de interpretação: {_band_label(prob)}",
    ]
    if pmi is not None:
        metrics.append(f"PMI: {float(pmi):.3f}")
    if interval and interval != "—":
        metrics.append(f"Intervalo entre árvores: {interval}")
    if q is not None:
        metrics.append(f"Carga do peptídeo (q): {float(q):.1f}")
    raw = res.get("pred_high_activity_prob_raw")
    if raw is not None:
        metrics.append(f"Probabilidade bruta (antes da calibração): {float(raw):.1%}")

    report.sections.append(
        ReportSection(title="2. Resultados numéricos", bullets=metrics)
    )

    if narrative:
        report.sections.append(
            ReportSection(
                title="3. Explicação em português",
                paragraphs=[narrative.strip()],
            )
        )

    if neighbors:
        headers = ["ID", "Nome", "Identidade", "Score", "MIC mediana (µM)"]
        rows = []
        for n in neighbors[:8]:
            ident = n.get("identity")
            rows.append(
                [
                    str(n.get("peptide_id") or "—"),
                    str(n.get("name") or "—")[:28],
                    f"{100 * float(ident):.0f}%" if ident is not None else "—",
                    f"{float(n['neighbor_score']):.3f}" if n.get("neighbor_score") is not None else "—",
                    str(n.get("mic_median_uM") if n.get("mic_median_uM") is not None else "—"),
                ]
            )
        report.sections.append(
            ReportSection(
                title="4. Vizinhos no treino",
                paragraphs=[
                    "Peptídeos parecidos ajudam a contextualizar a predição "
                    "(identidade de sequência e atividade observada)."
                ],
                table_headers=headers,
                table_rows=rows,
            )
        )

    if shap_top:
        headers = ["Descritor", "Valor SHAP", "Efeito"]
        rows = []
        for row in shap_top[:10]:
            val = row.get("shap_value")
            if val is None:
                continue
            efeito = "favorece alta atividade" if float(val) >= 0 else "desfavorece"
            rows.append(
                [
                    str(row.get("label") or row.get("feature") or "—"),
                    f"{float(val):+.4f}",
                    efeito,
                ]
            )
        report.sections.append(
            ReportSection(
                title="5. Contribuições SHAP (locais)",
                paragraphs=[
                    "Valores positivos empurram para alta atividade; negativos, para o contrário. "
                    "É explicação do modelo, não prova biológica."
                ],
                table_headers=headers,
                table_rows=rows,
            )
        )

    report.sections.append(
        ReportSection(
            title="6. Notas de uso",
            bullets=[
                "Alta atividade no modelo = MIC ≤ 3,4 µM no treino.",
                "Use o resultado para priorizar ensaios in vitro.",
                "Confirme sempre na bancada — o relatório não substitui experimento.",
                "Projeto: PepMem-AI · peçonha de Tityus stigmurus · InovAI Lab / UFRN.",
            ],
        )
    )
    return report


def build_batch_report(
    *,
    target_label: str,
    rows: list[dict[str, Any]],
    narrative: str | None = None,
    charge_note: str = "",
) -> PepMemReport:
    """Monta relatório de lote FASTA."""
    now = datetime.now().strftime("%d/%m/%Y %H:%M")
    ok = [r for r in rows if r.get("prob_calibrada") is not None]
    n_high = sum(1 for r in ok if float(r["prob_calibrada"]) >= 0.70)

    report = PepMemReport(
        title="Relatório PepMem-AI — Lote",
        subtitle="Predição em lote · InovAI Lab / UFRN",
        generated_at=now,
    )
    report.sections.append(
        ReportSection(
            title="1. Resumo do lote",
            bullets=[
                f"Membrana-alvo: {target_label}",
                f"Peptídeos no lote: {len(rows)}",
                f"Predições válidas: {len(ok)}",
                f"Com ≥ 70% (fortes): {n_high}",
                f"Carga: {charge_note or 'estimada por sequência'}",
                f"Gerado em: {now}",
            ],
        )
    )
    if narrative:
        report.sections.append(
            ReportSection(
                title="2. Explicação em português",
                paragraphs=[narrative.strip()],
            )
        )

    headers = ["Cabeçalho", "Sequência", "Prob.", "PMI", "Faixa", "Carga"]
    table = []
    ranked = sorted(
        ok,
        key=lambda r: float(r["prob_calibrada"]),
        reverse=True,
    )
    for r in ranked:
        table.append(
            [
                str(r.get("header") or "—")[:36],
                str(r.get("sequence") or "—")[:22],
                f"{float(r['prob_calibrada']):.1%}",
                str(r.get("pmi") if r.get("pmi") is not None else "—"),
                str(r.get("faixa") or "—")[:28],
                str(r.get("q_peptide") if r.get("q_peptide") is not None else "—"),
            ]
        )
    report.sections.append(
        ReportSection(
            title="3. Resultados ordenados",
            paragraphs=["Ordenado por probabilidade calibrada (maior → menor)."],
            table_headers=headers,
            table_rows=table,
        )
    )
    report.sections.append(
        ReportSection(
            title="4. Notas de uso",
            bullets=[
                "Priorize os de ≥ 70% para ensaio, olhando também PMI e faixa.",
                "Confirme na bancada — priorização, não substituto experimental.",
                "PepMem-AI · InovAI Lab / UFRN.",
            ],
        )
    )
    return report


def build_shap_overview_report(
    *,
    n_train: int,
    narrative: str | None = None,
    baseline_importance: list[dict[str, Any]] | None = None,
    multimodal_importance: list[dict[str, Any]] | None = None,
) -> PepMemReport:
    """Monta relatório do panorama SHAP global."""
    now = datetime.now().strftime("%d/%m/%Y %H:%M")
    report = PepMemReport(
        title="Relatório PepMem-AI — SHAP global",
        subtitle="Importância dos descritores · InovAI Lab / UFRN",
        generated_at=now,
    )
    report.sections.append(
        ReportSection(
            title="1. Contexto",
            bullets=[
                f"Pares MIC no treino: {n_train}",
                "Rótulo: MIC ≤ 3,4 µM = alta atividade",
                f"Gerado em: {now}",
            ],
        )
    )
    if narrative:
        report.sections.append(
            ReportSection(
                title="2. Explicação em português",
                paragraphs=[narrative.strip()],
            )
        )

    def _imp_table(rows: list[dict[str, Any]] | None) -> tuple[list[str], list[list[str]]]:
        headers = ["Descritor", "Média |SHAP|", "Grupo"]
        out: list[list[str]] = []
        if not rows:
            return headers, out
        ranked = sorted(
            rows,
            key=lambda r: float(r.get("mean_abs_shap") or abs(float(r.get("shap_value") or 0))),
            reverse=True,
        )
        for r in ranked[:12]:
            out.append(
                [
                    str(r.get("label") or r.get("feature") or "—"),
                    f"{float(r.get('mean_abs_shap') or r.get('shap_value') or 0):.4f}",
                    str(r.get("group") or "—"),
                ]
            )
        return headers, out

    h, rows = _imp_table(baseline_importance)
    report.sections.append(
        ReportSection(
            title="3. Importância — baseline",
            table_headers=h,
            table_rows=rows,
        )
    )
    h2, rows2 = _imp_table(multimodal_importance)
    report.sections.append(
        ReportSection(
            title="4. Importância — multimodal",
            table_headers=h2,
            table_rows=rows2 or [["—", "—", "Relatório multimodal ausente neste deploy"]],
        )
    )
    report.sections.append(
        ReportSection(
            title="5. Notas de uso",
            bullets=[
                "SHAP explica o modelo; não é prova biológica isolada.",
                "Combine com PMI, probabilidade calibrada e ensaios in vitro.",
                "PepMem-AI · InovAI Lab / UFRN.",
            ],
        )
    )
    return report


def build_ranking_report(
    *,
    sequence: str,
    lambda_tox: float,
    rows: list[dict[str, Any]],
    narrative: str | None = None,
    type_filter: list[str] | None = None,
) -> PepMemReport:
    """Monta relatório do ranking multi-alvo."""
    now = datetime.now().strftime("%d/%m/%Y %H:%M")
    report = PepMemReport(
        title="Relatório PepMem-AI — Ranking",
        subtitle="Priorização multi-alvo · InovAI Lab / UFRN",
        generated_at=now,
    )
    filt = ", ".join(type_filter) if type_filter else "todos os tipos"
    report.sections.append(
        ReportSection(
            title="1. Identificação",
            bullets=[
                f"Sequência: {sequence}",
                f"Comprimento: {len(sequence)} aminoácidos",
                f"Penalização toxicidade (λ): {lambda_tox:.2f}",
                f"Filtro de tipo de alvo: {filt}",
                f"Alvos no ranking: {len(rows)}",
                f"Gerado em: {now}",
            ],
        )
    )
    if narrative:
        report.sections.append(
            ReportSection(
                title="2. Explicação em português",
                paragraphs=[narrative.strip()],
            )
        )

    ranked = sorted(
        [r for r in rows if r.get("final_score") is not None],
        key=lambda r: float(r["final_score"]),
        reverse=True,
    )
    headers = ["Alvo", "Tipo", "PMI", "PMI sel.", "Prob.", "Score final"]
    table: list[list[str]] = []
    for r in ranked:
        prob = r.get("pred_high_activity_prob")
        if isinstance(prob, str):
            prob_s = prob
        elif prob is not None:
            prob_s = f"{float(prob):.1%}"
        else:
            prob_s = "—"
        table.append(
            [
                str(r.get("target_id") or "—")[:36],
                str(r.get("target_type") or "—")[:16],
                f"{float(r['pmi']):.3f}" if r.get("pmi") is not None else "—",
                f"{float(r['pmi_sel']):.3f}" if r.get("pmi_sel") is not None else "—",
                prob_s,
                f"{float(r['final_score']):.4f}",
            ]
        )
    report.sections.append(
        ReportSection(
            title="3. Matriz ordenada por score final",
            paragraphs=[
                "Score ≈ probabilidade de alta atividade − λ × toxicidade estimada "
                "+ bônus de PMI seletivo. Topo = priorizar primeiro na bancada."
            ],
            table_headers=headers,
            table_rows=table,
        )
    )
    report.sections.append(
        ReportSection(
            title="4. Notas de uso",
            bullets=[
                "Alta atividade no modelo = MIC ≤ 3,4 µM.",
                "λ alto = mais cautela com toxicidade proxy (célula normal).",
                "Confirme sempre in vitro — priorização, não substituto experimental.",
                "PepMem-AI · InovAI Lab / UFRN.",
            ],
        )
    )
    return report


# --- renderers ---


def report_to_markdown(report: PepMemReport) -> str:
    """Gera Markdown limpo e organizado."""
    lines = [
        f"# {report.title}",
        "",
        f"*{report.subtitle}*",
        "",
        f"**Gerado em:** {report.generated_at}",
        "",
        "---",
        "",
    ]
    for sec in report.sections:
        lines.append(f"## {sec.title}")
        lines.append("")
        for p in sec.paragraphs:
            lines.append(p)
            lines.append("")
        if sec.bullets:
            for b in sec.bullets:
                lines.append(f"- {b}")
            lines.append("")
        if sec.table_headers and sec.table_rows is not None:
            hdr = sec.table_headers
            lines.append("| " + " | ".join(hdr) + " |")
            lines.append("| " + " | ".join("---" for _ in hdr) + " |")
            for row in sec.table_rows:
                cells = [c.replace("|", "/") for c in row]
                # pad
                while len(cells) < len(hdr):
                    cells.append("—")
                lines.append("| " + " | ".join(cells[: len(hdr)]) + " |")
            lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("*PepMem-AI — priorização in vitro · InovAI Lab / UFRN*")
    lines.append("")
    return "\n".join(lines)


def report_to_docx_bytes(report: PepMemReport) -> bytes:
    """Gera DOCX em memória."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_heading(report.title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph(report.subtitle)
    for run in sub.runs:
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    meta = doc.add_paragraph(f"Gerado em: {report.generated_at}")
    for run in meta.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for sec in report.sections:
        doc.add_heading(sec.title, level=1)
        for p in sec.paragraphs:
            doc.add_paragraph(p)
        for b in sec.bullets:
            doc.add_paragraph(b, style="List Bullet")
        if sec.table_headers and sec.table_rows:
            ncols = len(sec.table_headers)
            table = doc.add_table(rows=1 + len(sec.table_rows), cols=ncols)
            table.style = "Table Grid"
            for j, name in enumerate(sec.table_headers):
                cell = table.rows[0].cells[j]
                cell.text = name
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for i, row in enumerate(sec.table_rows):
                for j in range(ncols):
                    table.rows[i + 1].cells[j].text = row[j] if j < len(row) else "—"

    foot = doc.add_paragraph()
    run = foot.add_run("PepMem-AI — priorização in vitro · InovAI Lab / UFRN")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def report_to_pdf_bytes(report: PepMemReport) -> bytes:
    """Gera PDF em memória (fpdf2)."""
    from fpdf import FPDF

    class ReportPDF(FPDF):
        def footer(self):
            self.set_y(-12)
            self.set_font("Helvetica", size=8)
            self.set_text_color(120, 120, 120)
            self.cell(0, 8, f"{self.page_no()} · PepMem-AI · InovAI Lab / UFRN", align="C")

    pdf = ReportPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=16)
    pdf.add_page()
    pdf.set_margins(18, 16, 18)

    def _safe(text: str) -> str:
        # Helvetica core font: evita caracteres fora de latin-1
        return (
            text.replace("—", "-")
            .replace("–", "-")
            .replace("×", "x")
            .replace("µ", "u")
            .replace("≥", ">=")
            .replace("≤", "<=")
            .replace("·", "|")
            .encode("latin-1", errors="replace")
            .decode("latin-1")
        )

    usable = pdf.epw

    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(40, 30, 50)
    pdf.multi_cell(usable, 8, _safe(report.title))
    pdf.ln(1)
    pdf.set_font("Helvetica", "I", 10)
    pdf.set_text_color(90, 90, 90)
    pdf.multi_cell(usable, 5, _safe(report.subtitle))
    pdf.set_font("Helvetica", size=9)
    pdf.multi_cell(usable, 5, _safe(f"Gerado em: {report.generated_at}"))
    pdf.ln(3)
    pdf.set_draw_color(180, 180, 180)
    y = pdf.get_y()
    pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
    pdf.ln(4)

    for sec in report.sections:
        pdf.set_x(pdf.l_margin)
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(40, 30, 50)
        pdf.multi_cell(usable, 7, _safe(sec.title))
        pdf.ln(1)
        pdf.set_text_color(30, 30, 30)
        pdf.set_font("Helvetica", size=10)
        for p in sec.paragraphs:
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(usable, 5, _safe(p))
            pdf.ln(1)
        for b in sec.bullets:
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(usable, 5, _safe(f"- {b}"))
        if sec.table_headers and sec.table_rows:
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 8)
            hdr_line = " | ".join(sec.table_headers)
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(usable, 4, _safe(hdr_line))
            pdf.set_font("Helvetica", size=8)
            for row in sec.table_rows:
                pdf.set_x(pdf.l_margin)
                line = " | ".join(row)
                if len(line) > 140:
                    line = line[:137] + "..."
                pdf.multi_cell(usable, 4, _safe(line))
        pdf.ln(3)

    out = pdf.output()
    if isinstance(out, (bytes, bytearray)):
        return bytes(out)
    return str(out).encode("latin-1", errors="replace")


def export_report_bundle(report: PepMemReport) -> dict[str, bytes]:
    """Retorna dict com md/docx/pdf prontos para download."""
    md = report_to_markdown(report).encode("utf-8")
    docx = report_to_docx_bytes(report)
    pdf = report_to_pdf_bytes(report)
    return {"md": md, "docx": docx, "pdf": pdf}
